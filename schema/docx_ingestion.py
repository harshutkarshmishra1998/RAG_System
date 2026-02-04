from __future__ import annotations

import os
import shutil
from pathlib import Path
from typing import List

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image
import pytesseract
import io

from schema.ingestion import (
    IngestedDocument,
    ContentBlock,
    ContentType,
    SourceMetadata,
    SourceType,
    BlockMetadata,
)

# ============================================================
# Runtime dependency resolution (executed at import time)
# Assumes `.env` already loaded
# ============================================================

def _resolve_tesseract() -> None:
    tesseract_path = os.getenv("TESSERACT_PATH")

    if tesseract_path:
        if not os.path.isfile(tesseract_path):
            raise RuntimeError(f"Invalid TESSERACT_PATH: {tesseract_path}")
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        return

    if shutil.which("tesseract"):
        return

    raise RuntimeError(
        "Tesseract not found. Set TESSERACT_PATH or add it to PATH."
    )


_resolve_tesseract()

# ============================================================
# DOCX ingestion
# ============================================================

def ingest_docx(docx_path: str | Path) -> IngestedDocument:
    """
    Ingest a DOCX file into an IngestedDocument.

    Handles:
    - Native text (paragraphs)
    - Tables
    - Embedded images (OCR)
    """

    docx_path = Path(docx_path)

    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX not found: {docx_path}")

    blocks: List[ContentBlock] = []

    document = Document(docx_path) # type: ignore

    # --------------------------------------------------------
    # Pass 1: Paragraph text
    # --------------------------------------------------------
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(
                ContentBlock(
                    content_type=ContentType.TEXT,
                    text=text,
                    metadata=BlockMetadata(
                        extra={"style": para.style.name} # type: ignore
                    ),
                )
            )

    # --------------------------------------------------------
    # Pass 2: Tables
    # --------------------------------------------------------
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append("\t".join(cell.text.strip() for cell in row.cells))

        table_text = "\n".join(rows).strip()
        if table_text:
            blocks.append(
                ContentBlock(
                    content_type=ContentType.TABLE,
                    text=table_text,
                    metadata=BlockMetadata(),
                )
            )

    # --------------------------------------------------------
    # Pass 3: Embedded images → OCR
    # --------------------------------------------------------
    for rel in document.part.rels.values():
        if rel.reltype == RT.IMAGE:
            image_bytes = rel.target_part.blob
            image = Image.open(io.BytesIO(image_bytes))

            ocr_text = pytesseract.image_to_string(image).strip()

            if ocr_text:
                blocks.append(
                    ContentBlock(
                        content_type=ContentType.IMAGE_OCR,
                        text=ocr_text,
                        metadata=BlockMetadata(
                            extra={"source": "docx_embedded_image"}
                        ),
                    )
                )

    return IngestedDocument(
        source=SourceMetadata(
            source_type=SourceType.DOCX,
            source_uri=str(docx_path),
            file_name=docx_path.name,
        ),
        blocks=blocks,
    )